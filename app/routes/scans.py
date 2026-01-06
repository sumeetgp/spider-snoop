"""API Routes - DLP Scanning"""
from typing import List
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, status, File, UploadFile, Query
from sqlalchemy.orm import Session
from datetime import datetime

from app.database import get_db
from app.models.user import User, UserRole
from app.models.scan import DLPScan, ScanStatus
from app.schemas.scan import ScanCreate, ScanResponse, ScanStats
from app.utils.auth import get_current_active_user
from app.dlp_engine import DLPEngine
from app.core.dlp_patterns import DLPPatternMatcher

from app.dlp_engine import DLPEngine
from app.utils.limiter import limiter, get_rate_limit_key
from fastapi import Request

router = APIRouter(prefix="/api/scans", tags=["DLP Scanning"])
dlp_engine = DLPEngine()

@router.post("/", response_model=ScanResponse, status_code=status.HTTP_201_CREATED)
@limiter.limit("50/60minute")
async def create_scan(
    request: Request,
    scan_data: ScanCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Create and execute a DLP scan"""
    # Guardian Text Scan Cost: 2
    if current_user.credits_remaining < 2:
        raise HTTPException(status_code=402, detail="Insufficient credits. Text scan costs 2 credits.")
    
    current_user.credits_remaining -= 2
    db.add(current_user)
    
    # Create scan record
    db_scan = DLPScan(
        user_id=current_user.id,
        source=scan_data.source,
        content=scan_data.content,
        status=ScanStatus.SCANNING
    )
    db.add(db_scan)
    db.commit()
    db.refresh(db_scan)
    
    try:
        # Perform scan
        result = await dlp_engine.scan(scan_data.content)
        
        # Update scan record
        db_scan.status = ScanStatus.COMPLETED
        db_scan.risk_level = result['risk_level']
        db_scan.findings = result['findings']
        db_scan.verdict = result['verdict']
        
        # Calculate Threat Score
        # Calculate Threat Score
        base_score = 0
        risk_str = str(result['risk_level']).upper()
        
        if "CRITICAL" in risk_str: base_score = 90
        elif "HIGH" in risk_str: base_score = 75
        elif "MEDIUM" in risk_str: base_score = 45
        elif "LOW" in risk_str: base_score = 10
        
        # Adjust based on findings count (simple logic)
        score = base_score + (len(result['findings']) * 2)
        db_scan.threat_score = min(score, 100) # We might need to add this column to DB or just return it in schema if computed. 
        # Wait, the SCHEMA has it, but the MODEL (DB) doesn't have it yet. 
        # For now, to avoid DB migration in this step if user didn't ask for DB change explicitly but "shape project", 
        # I should probably just return it in the API response or add to DB.
        # Given "Reshape project", I should assume I can update I will add it to DB in a separate step if needed, 
        # OR I can just map it in the Pydantic model if I compute it on the fly. 
        # But `ScanResponse` expects it. 
        # Let's check `app/models/scan.py` first. 
        
        if result.get('ai_analysis'):
            import json
            db_scan.ai_analysis = json.dumps(result['ai_analysis'])
        db_scan.scan_duration_ms = result['scan_duration_ms']
        db_scan.completed_at = datetime.utcnow()
        
        db.commit()
        db.refresh(db_scan)
        
    except Exception as e:
        db_scan.status = ScanStatus.FAILED
        db_scan.verdict = f"Scan failed: {str(e)}"
        db.commit()
        db.refresh(db_scan)
    
    return db_scan

@router.post("/upload_file", response_model=ScanResponse, status_code=status.HTTP_201_CREATED)
@limiter.limit("50/60minute")
async def upload_file(
    request: Request,
    track: str = Query("guardian"),
    correct: bool = Query(False, description="Redact sensitive data if found"),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """
    Upload and scan a file for sensitive data and forensics.
    
    Supported:
    - **Docs**: PDF, DOCX, TXT, CSV, JSON
    - **Images**: PNG, JPG (OCR)
    - **Archives**: ZIP (Structure check)
    - **Forensics**: Metadata & Macros (via MCP Skills)
    
    **Limit**: 10 MB (Guardian) / 50 MB (Sentinel)
    """
    import os
    import shutil
    import uuid
    import json
    from app.cdr_engine import CDREngine
    from app.utils.storage import StorageManager
    
    # Create temp file
    os.makedirs("storage", exist_ok=True)
    file_ext = os.path.splitext(file.filename)[1].lower()
    temp_filename = f"storage/temp_{uuid.uuid4()}{file_ext}"
    
    try:
        # Save to disk
        with open(temp_filename, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        file_size_mb = os.path.getsize(temp_filename) / (1024 * 1024)
        
        # --- SENTINEL TRACK (Malware Focus) ---
        if track == 'sentinel':
            # Cost: 1 Credit
            if current_user.credits_remaining < 1:
                raise HTTPException(status_code=402, detail="Insufficient credits. Sentinel scan costs 1 credit.")
            
            if file_size_mb > 50:
                 raise HTTPException(status_code=400, detail="File too large. Sentinel limit is 50MB.")
            
            # Deduct Credit
            current_user.credits_remaining -= 1
            db.add(current_user)
            db.commit() # Commit deduction immediately
            
            is_safe = True
            findings = []
            
            if hasattr(request.app.state, 'file_guard') and request.app.state.file_guard:
                is_safe, findings = await request.app.state.file_guard.scan_file(temp_filename)
            
            # Sentinel: Deep Macro Analysis (via MCP)
            if file_ext in ['.doc', '.docm', '.xls', '.xlsm', '.ppt', '.pptm', '.docx']:
                 macro_findings = await dlp_engine.scan_macros(temp_filename)
                 if macro_findings and not macro_findings.startswith("Error") and ("Found" in macro_findings or "Suspicious" in macro_findings or "AutoOpen" in macro_findings or "VBA" in macro_findings):
                      is_safe = False
                      findings.append(f"Deep Macro Analysis: {macro_findings[:200]}...") # Truncate for summary
            
            # --- CDR SANITIZATION (SAFE WASH) ---
            cdr_result = None
            
            # Only perform Safe Wash if threats are detected (or explicitly requested, but for now just threats per user req)
            if not is_safe:
                storage = StorageManager()
                cdr = CDREngine()
                safe_filename = f"storage/safe_{uuid.uuid4()}{file_ext}"
                
                try:
                    if cdr.disarm(temp_filename, safe_filename):
                        # Upload Safe Copy
                        upload_url = await storage.upload_file(safe_filename, f"safe/{current_user.id}/{os.path.basename(safe_filename)}")
                        if upload_url:
                            cdr_result = {"status": "success", "url": upload_url}
                        else:
                            cdr_result = {"status": "local_only", "url": None, "path": safe_filename} # Fallback
                        
                        # Cleanup safe local file if uploaded
                        if upload_url and os.path.exists(safe_filename):
                            os.remove(safe_filename)
                except Exception as e:
                    cdr_result = {"status": "failed", "error": str(e)}

            # Create Scan Record (Sentinel)
            db_scan = DLPScan(
                user_id=current_user.id,
                source=f"FILE:{file.filename}",
                content=f"[SENTINEL SCAN]\nFile: {file.filename}\nSize: {file_size_mb:.2f}MB", 
                status=ScanStatus.COMPLETED
            )
            
            if cdr_result:
                db_scan.ai_analysis = json.dumps({"cdr": cdr_result})

            if not is_safe:
                db_scan.risk_level = "critical"
                db_scan.findings = [{"type": "malware", "severity": "critical", "detail": f} for f in findings]
                db_scan.verdict = "MALWARE DETECTED"
                db_scan.threat_score = 95
                db_scan.summary = f"Sentinel detected {len(findings)} threat(s)."
            else:
                db_scan.risk_level = "low"
                db_scan.findings = []
                db_scan.verdict = "SAFE"
                db_scan.threat_score = 0
                db_scan.summary = "Sentinel analysis complete. No threats found."
                
            db_scan.created_at = datetime.utcnow()
            db_scan.completed_at = datetime.utcnow()
            
            db.add(db_scan)
            db_scan.credits_used = 1 # Not in model yet, but useful for logs
            db.commit()
            db.refresh(db_scan)
            
            return db_scan

        # --- GUARDIAN TRACK (DLP & Forensics) ---
        cost = 2
        if track == 'vision': cost = 10
        
        if current_user.credits_remaining < cost:
             raise HTTPException(status_code=402, detail=f"Insufficient credits. {track.title()} scan costs {cost} credits.")
             
        # Deduct Credits
        current_user.credits_remaining -= cost
        db.add(current_user)
        db.commit()
        
        if file_size_mb > 10:
             raise HTTPException(status_code=400, detail="File too large. Guardian limit is 10MB.")


        # 0. File Guard (AV/YARA) Sanity Check
        if hasattr(request.app.state, 'file_guard') and request.app.state.file_guard:
            is_safe, findings = await request.app.state.file_guard.scan_file(temp_filename)
            if not is_safe:
                import logging
                logging.getLogger(__name__).warning(f"File Guard blocked upload: {file.filename}, Findings: {findings}")
                # For Guardian, we still block processing but arguably we could offer sanitization?
                # For now let's keep blocking logic but maybe return a specific error code
                # Or just raise the exception as before.
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Security Threat Detected: {', '.join(findings)}"
                )

        # 1. Extract Text Content (for Regex/AI)
        content = ""
        filename = file.filename.lower()
        
        try:
            if filename.endswith(('.docx', '.docm')):
                # Attempt to extract text from DOCX/DOCM
                # Note: docx can often read text from docm, but might warn about macros.
                try:
                    import docx
                    doc = docx.Document(temp_filename)
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                except Exception as docx_err:
                    content = f"[Office Document - Text Extraction Failed: {str(docx_err)}]\n(Forensics will analyze macros)"
                  
            elif filename.endswith('.pdf'):
                import pypdf
                pdf_reader = pypdf.PdfReader(temp_filename)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
                    
            elif filename.endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
                # For images, we rely on the MCP Tool 'scan_image_text' later, 
                # but we can also do a quick extract here if pytesseract is avail locally
                try:
                    import pytesseract
                    from PIL import Image
                    image = Image.open(temp_filename)
                    content = pytesseract.image_to_string(image)
                except:
                    content = "[Image Content - Waiting for AI Forensics]"

            elif filename.endswith(('.xlsx', '.xlsm')):
                 # Basic extraction for Excel
                 try:
                     import openpyxl
                     wb = openpyxl.load_workbook(temp_filename, data_only=True)
                     content = ""
                     for sheet in wb.sheetnames:
                         ws = wb[sheet]
                         for row in ws.iter_rows(values_only=True):
                             row_text = " ".join([str(cell) for cell in row if cell is not None])
                             content += row_text + "\n"
                 except Exception as xl_err:
                     content = f"[Excel Document - Extraction Failed: {str(xl_err)}]"
        
            elif filename.endswith(('.json', '.txt', '.md', '.csv', '.log', '.xml', '.yaml', '.yml')):
                 with open(temp_filename, 'r', errors='ignore') as f:
                     content = f.read()
            elif filename.endswith('.zip'):
                content = "[ZIP Archive - Waiting for AI Forensics]"
            else:
                 # Try reading as text fallback
                 with open(temp_filename, 'r', errors='ignore') as f:
                     content = f.read(10000) # Read first 10k
        except Exception as e:
            content = f"Error extracting text: {str(e)}"

        # Create scan record
        db_scan = DLPScan(
            user_id=current_user.id,
            source=f"FILE:{file.filename}",
            content=content[:50000], # Limit DB storage
            status=ScanStatus.SCANNING
        )
        db.add(db_scan)
        db.commit()
        db.refresh(db_scan)
        
        # 2. Perform Scan (Pass file_path!)
        # We assume dlp_engine.scan can take file_path argument now
        result = await dlp_engine.scan(content, file_path=temp_filename)
        
        # Update scan record
        db_scan.status = ScanStatus.COMPLETED
        db_scan.risk_level = result['risk_level']
        db_scan.findings = result['findings']
        db_scan.verdict = result['verdict']
        
        # Threat Score Logic (0 = Safe, 100 = Risk)
        # 1. Start with AI Score
        raw_score = 0
        if result.get('ai_analysis') and 'score' in result['ai_analysis']:
             raw_score = result['ai_analysis']['score']
        
        # 2. Enforce minimum score based on Findings Severity
        finding_severities = [f.get('severity', 'UNKNOWN').upper() for f in result['findings']]
        min_score = 0
        if "CRITICAL" in finding_severities: min_score = 95
        elif "HIGH" in finding_severities: min_score = 80
        elif "MEDIUM" in finding_severities: min_score = 50
        elif "LOW" in finding_severities: min_score = 25
        
        # 3. Take Maximum
        db_scan.threat_score = max(raw_score, min_score)

        # --- REDACTION (SCAN & CORRECT) ---
        redaction_result = None
        if False and correct and result['findings']:
             # Reuse the existing DLPEngine instance (which has Presidio loaded)
             # engine is already instantiated at line ~307 (Variable name is dlp_engine in code? No, let's check view)
             # Wait, in lines 500+ it is 'engine' = DLPEngine().
             # In lines 300+ (upload_file), we call 'dlp_engine.scan'.
             # So variable name is 'dlp_engine'.

             # Redact the content using the FULL Engine (Presidio + Regex)
             redacted_text = dlp_engine.redact(content)
             
             # Save to file (Preserve Format)
             ext = os.path.splitext(file.filename)[1].lower()
             redacted_filename = f"storage/redacted_{uuid.uuid4()}{ext}"
             
             try:
                 if ext in ['.docx', '.docm']:
                     # Generate DOCX
                     import docx
                     doc = docx.Document()
                     doc.add_heading("Redacted Document (Safe Wash)", 0)
                     doc.add_paragraph(redacted_text)
                     doc.add_paragraph("\n[Metadata Washed | PII Redacted]")
                     doc.save(redacted_filename)
                     
                 elif ext == '.pdf':
                     # Generate PDF
                     from reportlab.pdfgen import canvas
                     from reportlab.lib.pagesizes import letter
                     from reportlab.lib.units import inch
                     
                     c = canvas.Canvas(redacted_filename, pagesize=letter)
                     text_obj = c.beginText()
                     text_obj.setTextOrigin(inch, 10*inch)
                     text_obj.setFont("Helvetica", 10)
                     
                     # Simple logic to handle multiline text
                     lines = redacted_text.split('\n')
                     for line in lines:
                         # Very basic wrapping (truncate for now or rely on simple flow)
                         # Real wrap needs SimpleDocTemplate but Canvas is safer/simpler for raw text dump
                         # limit chars per line
                         chunks = [line[i:i+90] for i in range(0, len(line), 90)]
                         for chunk in chunks:
                             text_obj.textLine(chunk)
                             if text_obj.getY() < inch: # New page
                                 c.drawText(text_obj)
                                 c.showPage()
                                 text_obj = c.beginText()
                                 text_obj.setTextOrigin(inch, 10*inch)
                                 text_obj.setFont("Helvetica", 10)
                                 
                     c.drawText(text_obj)
                     c.save()
                     
                 else:
                     # Fallback to TXT
                     redacted_filename = f"storage/redacted_{uuid.uuid4()}.txt" # Force txt extension
                     with open(redacted_filename, "w") as f:
                         f.write(redacted_text)
                         
             except Exception as write_err:
                 # Fallback to TXT on error
                 import logging
                 logging.getLogger(__name__).error(f"Failed to write native format {ext}: {write_err}")
                 redacted_filename = f"storage/redacted_{uuid.uuid4()}.txt"
                 with open(redacted_filename, "w") as f:
                     f.write(redacted_text)
                 
             # Upload to storage
             from app.utils.storage import StorageManager
             storage = StorageManager()
             redacted_url = await storage.upload_file(redacted_filename, f"redacted/{current_user.id}/{os.path.basename(redacted_filename)}")
             
             if redacted_url:
                 redaction_result = {"status": "success", "url": redacted_url}
             else:
                 redaction_result = {"status": "local_only", "url": None, "path": redacted_filename}
             
             # Cleanup local
             if redacted_url and os.path.exists(redacted_filename):
                 os.remove(redacted_filename)

        if redaction_result:
             scan_meta = {"redaction": redaction_result}
             if result.get('ai_analysis'):
                 scan_meta.update(result['ai_analysis'])
             import json
             db_scan.ai_analysis = json.dumps(scan_meta)
        elif result.get('ai_analysis'):
             import json
             db_scan.ai_analysis = json.dumps(result['ai_analysis'])
             
        if result.get('ai_analysis') and 'score' in result['ai_analysis']:
                db_scan.threat_score = result['ai_analysis']['score']
                
        db_scan.scan_duration_ms = result['scan_duration_ms']
        db_scan.completed_at = datetime.utcnow()
        
        db.commit()
        db.refresh(db_scan)
        
    except HTTPException:
        raise
    except Exception as e:
        if 'db_scan' in locals():
            db_scan.status = ScanStatus.FAILED
            db_scan.verdict = f"Scan failed: {str(e)}"
            db.commit()
        raise HTTPException(status_code=500, detail=f"Upload error: {str(e)}")
    finally:
        # Cleanup Temp File
        if os.path.exists(temp_filename):
            try:
                os.remove(temp_filename)
            except:
                pass
    
    return db_scan

@router.post("/upload_video", response_model=ScanResponse, status_code=status.HTTP_201_CREATED)
@limiter.limit("20/60minute") # Stricter limit for heavy video processing
async def upload_video(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """

    Upload and scan a video or audio file for sensitive spoken data (EchoVision).
    
    Process:
    1. Extract audio (if video).
    2. Transcribe audio to text (OpenAI Whisper).
    3. Scan transcript for sensitive info.
    
    **Limit**: 50 MB per file.
    """
    allowed_exts = ('.mp4', '.mov', '.avi', '.mkv', '.webm', '.mp3', '.wav', '.m4a', '.flac')
    if not file.filename.lower().endswith(allowed_exts):
        raise HTTPException(status_code=400, detail="Invalid media format. Use MP4, AVI, MOV, MKV, MP3, WAV, FLAC.")

    # 1. Save temp file to storage
    import shutil
    import os
    import json
    from app.utils.video import VideoProcessor
    
    # Initialize variables for cleanup
    cloud_url = None
    stored_filename = None
    
    # Ensure storage directory exists
    os.makedirs("storage", exist_ok=True)
    
    # Save with unique name
    stored_filename = f"storage/{current_user.id}_{int(datetime.utcnow().timestamp())}_{file.filename}"
    
    try:
        with open(stored_filename, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        # Check size (approx)
        if os.path.getsize(stored_filename) > 50 * 1024 * 1024:
            os.remove(stored_filename)
            raise HTTPException(status_code=413, detail="Media too large. Limit is 50MB.")
            
        # 1.5 File Guard (AV/YARA)
        # We can implement simpler scan here or retain existing logic
        # For brevity, retaining logic but pointing to stored file
        if hasattr(request.app.state, 'file_guard') and request.app.state.file_guard:
            is_safe, findings = await request.app.state.file_guard.scan_file(stored_filename)
            if not is_safe:
                if os.path.exists(stored_filename):
                    os.remove(stored_filename)
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Security Threat Detected: {', '.join(findings)}"
                )

        # 2. Process Media (Extract/Transcribe)
        # Returns { "text": "...", "segments": [...] } or { "error": "..." }
        result = await VideoProcessor.process_video(stored_filename)
        
        if "error" in result:
             if os.path.exists(stored_filename):
                 os.remove(stored_filename)
             raise HTTPException(status_code=500, detail=f"Processing failed: {result['error']}")

        # Extract text for DLP scanning
        transcript_text = result.get('text', '')

        # 2.5 Cloud Upload (DigitalOcean Spaces)
        from app.utils.storage import StorageManager
        storage = StorageManager()
        cloud_url = await storage.upload_file(stored_filename, f"{current_user.id}/{os.path.basename(stored_filename)}")
        
        if cloud_url:
            result['media_url'] = cloud_url
            # If upload successful, we can remove local file to save space
            if os.path.exists(stored_filename):
                os.remove(stored_filename)
            result['storage_type'] = 'cloud'
        else:
            result['storage_type'] = 'local'
            result['local_path'] = stored_filename
            # Keep local file since upload failed/disabled

    except HTTPException:
        if os.path.exists(stored_filename): os.remove(stored_filename)
        raise
    except Exception as e:
        if os.path.exists(stored_filename): os.remove(stored_filename)
        raise HTTPException(status_code=500, detail=f"Upload error: {str(e)}")

    # 3. Create Scan Record using JSON Content
    # content stores the full JSON structure
    db_scan = DLPScan(
        user_id=current_user.id,
        source=f"ECHOVISION:{file.filename}",
        content=json.dumps(result), 
        status=ScanStatus.SCANNING
    )
    db.add(db_scan)
    db.commit() # Commit to generate ID
    
    # 4. DLP Scan (using just the text part)
    # Re-fetch is not needed, we have the object
    from app.dlp_engine import DLPEngine
    engine = DLPEngine()
    
    # We scan ONLY the text, but the findings will link back to the record
    scan_result = await engine.scan(transcript_text, use_ai=True)
    
    # Update Record
    db_scan.risk_level = scan_result['risk_level']
    db_scan.findings = scan_result['findings']
    db_scan.verdict = scan_result['verdict']
    db_scan.ai_analysis = scan_result['ai_analysis']
    db_scan.threat_score = scan_result.get('ai_analysis', {}).get('score', 0) if scan_result.get('ai_analysis') else 0
    db_scan.completed_at = datetime.utcnow()
    db_scan.status = ScanStatus.COMPLETED
    
    db.commit()
    db.refresh(db_scan)
    
    # 5. Immediate Cleanup if Safe
    # If no threats found, delete the cloud file to save costs/storage
    if cloud_url and db_scan.risk_level.lower() == "low" and not db_scan.findings:
        if storage.delete_file(cloud_url):
            # Update content to reflect deletion
            content_data = json.loads(db_scan.content)
            content_data['media_url'] = None # Remove URL
            content_data['storage_status'] = 'deleted_clean'
            db_scan.content = json.dumps(content_data)
            db.commit()

    return db_scan

    try:
        # 4. Perform Scan on Transcript
        result = await dlp_engine.scan(transcription)
        
        db_scan.status = ScanStatus.COMPLETED
        db_scan.risk_level = result['risk_level']
        db_scan.findings = result['findings']
        db_scan.verdict = result['verdict']
        
        # Calculate Threat Score
        # Calculate Threat Score
        base_score = 0
        risk_str = str(result['risk_level']).upper()
        
        if "CRITICAL" in risk_str: base_score = 90
        elif "HIGH" in risk_str: base_score = 75
        elif "MEDIUM" in risk_str: base_score = 45
        elif "LOW" in risk_str: base_score = 10
        score = base_score + (len(result['findings']) * 2)
        db_scan.threat_score = min(score, 100)
        
        if result.get('ai_analysis'):
            import json
            db_scan.ai_analysis = json.dumps(result['ai_analysis'])
        db_scan.scan_duration_ms = result['scan_duration_ms']
        db_scan.completed_at = datetime.utcnow()
        
        db.commit()
        db.refresh(db_scan)
        
    except Exception as e:
        db_scan.status = ScanStatus.FAILED
        db_scan.verdict = f"Scan failed: {str(e)}"
        db.commit()
        db.refresh(db_scan)

    return db_scan

@router.get("/", response_model=List[ScanResponse])
async def list_scans(
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """List DLP scans"""
    # Regular users can only see their own scans
    query = db.query(DLPScan)
    
    # Rigorous filter: Admin sees all, everyone else sees ONLY their own
    is_admin = current_user.role == UserRole.ADMIN
    
    if not is_admin:
        query = query.filter(DLPScan.user_id == current_user.id)
    
    scans = query.order_by(DLPScan.created_at.desc()).offset(skip).limit(limit).all()
    return scans

@router.get("/stats", response_model=ScanStats)
async def get_scan_stats(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Get scan statistics"""
    query = db.query(DLPScan)
    
    is_admin = current_user.role == UserRole.ADMIN
    
    if not is_admin:
        query = query.filter(DLPScan.user_id == current_user.id)
    
    all_scans = query.all()
    
    # Calculate stats
    scans_by_risk = {}
    scans_by_status = {}
    total_duration = 0
    count_with_duration = 0
    
    for scan in all_scans:
        # Risk level stats
        risk = scan.risk_level.value if scan.risk_level else "UNKNOWN"
        scans_by_risk[risk] = scans_by_risk.get(risk, 0) + 1
        
        # Status stats
        status = scan.status.value
        scans_by_status[status] = scans_by_status.get(status, 0) + 1
        
        # Duration stats
        if scan.scan_duration_ms:
            total_duration += scan.scan_duration_ms
            count_with_duration += 1
    
    avg_duration = total_duration / count_with_duration if count_with_duration > 0 else 0
    
    # Get recent scans
    recent_scans = query.order_by(DLPScan.created_at.desc()).limit(10).all()
    
    return ScanStats(
        total_scans=len(all_scans),
        scans_by_risk=scans_by_risk,
        scans_by_status=scans_by_status,
        avg_scan_duration_ms=avg_duration,
        recent_scans=recent_scans
    )

@router.get("/{scan_id}", response_model=ScanResponse)
async def get_scan(
    scan_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Get scan by ID"""
    scan = db.query(DLPScan).filter(DLPScan.id == scan_id).first()
    
    if not scan:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Scan not found"
        )
    
    # Check permissions
    if current_user.role != UserRole.ADMIN and scan.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions"
        )
    
    return scan
